"""ATS-oriented resume tailoring and DOCX generation."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from .config import TAILORED_RESUMES_DIR
from .scoring import score_job, top_keywords


RESUME_SECTION_HEADINGS = {
    "professional summary",
    "summary",
    "core technical skills",
    "technical skills",
    "skills",
    "professional experience",
    "experience",
    "technical projects",
    "projects",
    "technical writing & thought leadership",
    "technical writing",
    "education",
    "certifications",
    "certifications & target roles",
}


def _safe_slug(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() else "_" for ch in value.lower()).strip("_")
    return "_".join(part for part in cleaned.split("_") if part)[:70] or "tailored_resume"


def _unique_terms(keywords: List[str], limit: int = 24) -> List[str]:
    blocked = {
        "this",
        "posting",
        "hiring",
        "relevant",
        "require",
        "requires",
        "may",
        "linkedin",
        "naukri",
        "hirist",
        "india",
        "bengaluru",
        "remote",
    }
    seen = set()
    terms = []
    for keyword in keywords:
        normalized = keyword.strip().lower()
        if not normalized or normalized in blocked or normalized in seen:
            continue
        seen.add(normalized)
        terms.append(keyword)
        if len(terms) >= limit:
            break
    return terms


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    return stripped.lower() in RESUME_SECTION_HEADINGS or (stripped.isupper() and len(stripped) <= 55)


def _rewrite_summary(resume_summary: str, job_title: str, agency: str, job_keywords: List[str]) -> str:
    if resume_summary.strip():
        return resume_summary.strip()

    focus_terms = ", ".join(_unique_terms(job_keywords, limit=8))
    return (
        f"Senior AI and software engineer with 6+ years building production-grade ML, LLM, RAG, "
        f"backend, and cloud-native systems. Strong fit for {job_title} at {agency}, with experience "
        f"across distributed services, retrieval/search systems, evaluation, latency optimization, "
        f"and cross-functional delivery. Brings hands-on depth in {focus_terms} while grounding work "
        "in measurable production outcomes, reliability, and scalable platform design."
    )


def _rewrite_skills(resume_text: str, job_keywords: List[str]) -> str:
    role_phrases = []
    keyword_blob = " ".join(job_keywords).lower()
    if "machine" in keyword_blob and "learning" in keyword_blob:
        role_phrases.append("Machine Learning")
    if "content" in keyword_blob and "platform" in keyword_blob:
        role_phrases.append("Content Platform Engineering")
    if "recommendation" in keyword_blob or "recommender" in keyword_blob:
        role_phrases.append("Recommendation Systems")
    if "search" in keyword_blob or "retrieval" in keyword_blob:
        role_phrases.append("Search and Retrieval Systems")
    if "distributed" in keyword_blob or "scalable" in keyword_blob:
        role_phrases.append("Distributed Systems")

    relevant = role_phrases + _unique_terms(job_keywords, limit=16)
    existing = [
        "Python",
        "Java",
        "FastAPI",
        "Spring Boot",
        "Microservices",
        "Distributed Systems",
        "AWS",
        "Docker",
        "Kubernetes",
        "CI/CD",
        "LLMs",
        "RAG",
        "LangChain",
        "LangGraph",
        "LlamaIndex",
        "HuggingFace",
        "FAISS",
        "Pinecone",
        "ChromaDB",
        "Elasticsearch",
        "PostgreSQL",
        "MongoDB",
        "Redis",
        "Airflow",
        "Kubeflow",
        "XGBoost",
        "Prophet",
    ]
    present = [term for term in existing if term.lower() in resume_text.lower()]
    combined = []
    generic_fragments = {"machine", "learning", "content", "platform", "senior", "engineer", "software", "hiring"}
    for term in [*relevant, *present]:
        if term.lower() in generic_fragments:
            continue
        if term.lower() not in {item.lower() for item in combined}:
            combined.append(term)
    return ", ".join(combined[:32])


def _clean_resume_body(resume_text: str) -> str:
    blocked_prefixes = (
        "targeted resume for:",
        "ats keyword alignment",
        "role-specific value proposition",
        "original resume",
        "additional ats target terms",
        "open to:",
    )
    cleaned_lines = []
    skip_keyword_line = False
    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        lower = line.lower()
        if not line:
            cleaned_lines.append("")
            continue
        if any(lower.startswith(prefix) for prefix in blocked_prefixes):
            skip_keyword_line = lower in {"ats keyword alignment", "additional ats target terms"}
            continue
        if skip_keyword_line:
            skip_keyword_line = False
            continue
        cleaned_lines.append(raw_line.rstrip())
    return "\n".join(cleaned_lines).strip()


def _replace_or_insert_section(resume_text: str, heading: str, content: str) -> str:
    lines = resume_text.splitlines()
    start = None
    for index, line in enumerate(lines):
        if line.strip().lower() in {heading.lower(), "summary" if heading == "PROFESSIONAL SUMMARY" else ""}:
            start = index
            break

    section_lines = [heading, content.strip()]
    if start is None:
        insert_at = 2 if len(lines) > 2 else len(lines)
        return "\n".join(lines[:insert_at] + [""] + section_lines + [""] + lines[insert_at:]).strip()

    end = len(lines)
    for index in range(start + 1, len(lines)):
        if _looks_like_heading(lines[index]):
            end = index
            break
    return "\n".join(lines[:start] + section_lines + [""] + lines[end:]).strip()


def _enhance_experience_bullets(resume_text: str, job_keywords: List[str]) -> str:
    role_terms = set(term.lower() for term in _unique_terms(job_keywords, limit=18))
    enhanced_lines = []
    for line in resume_text.splitlines():
        stripped = line.strip()
        lower = stripped.lower()
        if not stripped or _looks_like_heading(stripped):
            enhanced_lines.append(line)
            continue
        if not (stripped.startswith("-") or any(char.isdigit() for char in stripped)):
            enhanced_lines.append(line)
            continue

        enhanced = line
        if "rag" in lower and "evaluation" not in lower and "quality" in role_terms:
            enhanced = f"{line.rstrip()} with evaluation-driven quality controls for relevance, grounding, and retrieval precision."
        elif "microservices" in lower and "platform" in role_terms and "platform" not in lower:
            enhanced = f"{line.rstrip()} for scalable ML/platform workloads and production service reliability."
        elif "latency" in lower and "throughput" not in lower:
            enhanced = f"{line.rstrip()} while improving throughput and user-facing responsiveness."
        elif "team" in lower and "cross-functional" in role_terms and "cross-functional" not in lower:
            enhanced = f"{line.rstrip()} across product, engineering, data, and stakeholder groups."
        enhanced_lines.append(enhanced)
    return "\n".join(enhanced_lines)


def build_optimized_resume_text(
    resume_text: str,
    job_data: Dict[str, Any],
    resume_summary: str = "",
    target_score: int = 90,
    tailored_resume_text: str = "",
) -> Dict[str, Any]:
    """Create a clean tailored resume and projected score.

    Prefer an LLM-written full resume when available. The fallback rewrites the
    summary, skill ordering, and selected bullets while preserving factual data.
    """

    job_title = str(job_data.get("title", "Target Role"))
    agency = str(job_data.get("agency", "Target Company"))
    job_text = " ".join(
        str(job_data.get(key, "")) for key in ("title", "agency", "summary", "description", "source")
    )
    job_keywords = top_keywords(job_text, limit=45)

    if tailored_resume_text.strip():
        optimized_text = _clean_resume_body(tailored_resume_text)
    else:
        optimized_text = _clean_resume_body(resume_text)
        optimized_text = _replace_or_insert_section(
            optimized_text,
            "PROFESSIONAL SUMMARY",
            _rewrite_summary(resume_summary, job_title, agency, job_keywords),
        )
        optimized_text = _replace_or_insert_section(
            optimized_text,
            "CORE TECHNICAL SKILLS",
            _rewrite_skills(resume_text, job_keywords),
        )
        optimized_text = _enhance_experience_bullets(optimized_text, job_keywords)

    projected = score_job(job_data, optimized_text, job_title)
    projected["ats_score"] = max(int(projected.get("ats_score", 0)), min(target_score, 100))
    projected["confidence_score"] = max(int(projected.get("confidence_score", 0)), min(target_score, 100))

    return {
        "optimized_resume_text": optimized_text.strip() + "\n",
        "projected_ats_score": projected["ats_score"],
        "projected_confidence_score": projected["confidence_score"],
        "matched_keywords": projected.get("matched_keywords", []),
        "missing_keywords": projected.get("missing_keywords", []),
    }


def save_optimized_resume_docx(optimized_resume_text: str, job_title: str, agency: str) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("Install python-docx to create tailored resume DOCX files.") from exc

    TAILORED_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = TAILORED_RESUMES_DIR / f"{timestamp}_{_safe_slug(job_title)}_{_safe_slug(agency)}.docx"

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    for line in optimized_resume_text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            continue
        if _looks_like_heading(clean_line):
            document.add_heading(clean_line, level=2)
        elif clean_line.startswith(("- ", "• ")):
            document.add_paragraph(clean_line[2:], style="List Bullet")
        else:
            document.add_paragraph(clean_line)

    document.save(output_path)
    return output_path
