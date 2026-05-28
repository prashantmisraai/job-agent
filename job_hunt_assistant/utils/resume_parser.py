"""Resume text extraction helpers for Streamlit uploads."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


def extract_resume_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if filename.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")

    if filename.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Install python-docx to upload .docx resumes: pip install python-docx") from exc

        document = Document(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    if filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Install pypdf to upload PDF resumes: pip install pypdf") from exc

        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    raise ValueError("Supported resume formats are .txt, .docx, and .pdf.")


def extract_resume_text_from_path(path: str | Path) -> str:
    file_path = Path(path)
    if file_path.suffix.lower() == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    if file_path.suffix.lower() == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Install python-docx to read .docx resumes: pip install python-docx") from exc

        document = Document(file_path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    raise ValueError("Supported local resume formats are .txt and .docx.")
